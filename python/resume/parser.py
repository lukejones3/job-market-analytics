"""Parse resume bytes into clean text. Supports .pdf, .docx, .txt.

Extracts text from all parts of a document:
- DOCX: paragraphs, tables (including nested), headers, footers, text boxes
- PDF: all pages including tables
- TXT: raw decoded text
"""

import io
import re
import logging
from typing import List

logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


def _clean_text(s: str) -> str:
    """Strip nbsp, collapse whitespace within lines, preserve newlines."""
    s = s.replace("\u00a0", " ")
    # Collapse multiple spaces/tabs but preserve newlines
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract ALL text from a DOCX file.

    python-docx's d.paragraphs only returns top-level paragraphs.
    Modern resume templates often use tables for multi-column layout,
    so we must walk tables, headers, and footers explicitly.
    """
    if docx is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")

    d = docx.Document(io.BytesIO(file_bytes))
    parts: List[str] = []

    # 1. Top-level body paragraphs
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)

    # 2. Text inside tables (recursive: tables can contain tables)
    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            parts.append(p.text)
                    # Nested tables inside cells
                    if cell.tables:
                        walk_tables(cell.tables)

    walk_tables(d.tables)

    # 3. Headers and footers (every section)
    for section in d.sections:
        for hp in section.header.paragraphs:
            if hp.text:
                parts.append(hp.text)
        for fp in section.footer.paragraphs:
            if fp.text:
                parts.append(fp.text)
        # Headers/footers can also contain tables
        if section.header.tables:
            walk_tables(section.header.tables)
        if section.footer.tables:
            walk_tables(section.footer.tables)

    # 4. Text boxes and shapes (rare but possible in resume templates)
    # python-docx exposes these via the underlying XML
    try:
        from docx.oxml.ns import qn
        body = d.element.body
        # Find all <w:t> elements anywhere in the document
        # This catches text inside text boxes, shapes, comments
        for t_elem in body.iter(qn("w:t")):
            txt = t_elem.text
            if txt and txt not in parts:
                parts.append(txt)
    except Exception as e:
        logger.warning("DOCX deep XML extraction failed: %s", e)

    return _clean_text("\n".join(parts))


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF, including text in tables."""
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed. pip install pdfplumber")

    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Standard text extraction
            t = page.extract_text() or ""
            if t:
                parts.append(t)
            # Tables (extract_text sometimes misses table cells)
            try:
                tables = page.extract_tables() or []
                for table in tables:
                    for row in table:
                        for cell in row:
                            if cell:
                                parts.append(str(cell))
            except Exception as e:
                logger.warning("PDF table extraction failed on page: %s", e)

    return _clean_text("\n".join(parts))


def _extract_txt_text(file_bytes: bytes) -> str:
    """Decode a plain text file."""
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        text = file_bytes.decode("latin-1", errors="ignore")
    return _clean_text(text)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Parse a resume file's bytes into clean text.

    Args:
        file_bytes: Raw bytes of the uploaded file
        filename: Filename (used for extension detection)

    Returns:
        Clean text content of the resume

    Raises:
        ValueError: If file type is unsupported
        RuntimeError: If required parser library is not installed
    """
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if suffix == "txt":
        return _extract_txt_text(file_bytes)
    if suffix == "pdf":
        return _extract_pdf_text(file_bytes)
    if suffix == "docx":
        return _extract_docx_text(file_bytes)

    raise ValueError(
        f"Unsupported file type: .{suffix}. Use .pdf, .docx, or .txt"
    )
